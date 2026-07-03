from __future__ import annotations

import re
from pathlib import Path

from .models import SampleRecord, new_id

_CHAR_MIN = 50


def _char_count(text: str) -> int:
    return len(re.sub(r"\s+", "", text))


def _label_from_path(path: Path) -> str:
    return path.stem


def _split_document_sections(text: str) -> dict[str, object]:
    """생기부 문서 본문에서 행발·세특·창체 구간을 나눕니다."""
    sections: dict[str, object] = {"행발": "", "세특": {}, "창체": {}}

    markers = [
        ("행발", r"행동특성\s*(?:및\s*종합의견)?"),
        ("세특", r"세부능력\s*(?:및\s*특기사항)?"),
        ("창체", r"창의적\s*체험활동"),
    ]

    hits: list[tuple[int, str]] = []
    for key, pattern in markers:
        match = re.search(pattern, text)
        if match:
            hits.append((match.start(), key))

    if not hits:
        if _char_count(text) >= _CHAR_MIN:
            sections["행발"] = text.strip()
        return sections

    hits.sort(key=lambda item: item[0])
    for idx, (start, key) in enumerate(hits):
        end = hits[idx + 1][0] if idx + 1 < len(hits) else len(text)
        chunk = text[start:end]
        chunk = re.sub(
            r"^.*?(?:행동특성|세부능력|창의적\s*체험활동).*?\n",
            "",
            chunk,
            count=1,
            flags=re.S,
        )
        chunk = chunk.strip()
        if not chunk or _char_count(chunk) < _CHAR_MIN:
            continue
        if key == "행발":
            sections["행발"] = chunk
        elif key == "창체":
            cast = sections["창체"]
            assert isinstance(cast, dict)
            cast["자율"] = chunk
        else:
            cast = sections["세특"]
            assert isinstance(cast, dict)
            cast["과목미상"] = chunk

    return sections


def _records_from_document_text(path: Path, text: str) -> list[SampleRecord]:
    text = text.strip()
    if _char_count(text) < _CHAR_MIN:
        raise ValueError(f"문서에서 충분한 텍스트를 읽지 못했습니다: {path.name}")

    sections = _split_document_sections(text)
    has_content = bool(sections.get("행발")) or bool(sections.get("세특")) or bool(sections.get("창체"))
    if not has_content:
        sections = {"행발": text}

    return [
        SampleRecord(
            id=new_id("sample"),
            label=_label_from_path(path),
            sections=sections,
            source_file=str(path),
        )
    ]


def _xlsx_to_text(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            cells = ["" if cell is None else str(cell).strip() for cell in row]
            line = "\t".join(cell for cell in cells if cell)
            if line:
                lines.append(line)
    wb.close()
    return "\n".join(lines)


def _parse_xlsx_rows(path: Path) -> list[SampleRecord]:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    records: list[SampleRecord] = []
    label_base = _label_from_path(path)

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            vals = ["" if cell is None else str(cell).strip() for cell in row]
            if len(vals) < 3:
                continue

            content = ""
            name = ""
            if vals[0].isdigit() and vals[1] and not vals[1].isdigit() and len(vals[1]) <= 6:
                name = vals[1]
                content = max(vals[2:], key=_char_count, default="")
            elif len(vals) > 3 and vals[0].isdigit() and vals[1].isdigit() and vals[2]:
                name = vals[2]
                content = vals[3] if len(vals) > 3 else ""

            if not name or _char_count(content) < _CHAR_MIN:
                continue

            subject = label_base.split("(")[0] if "(" in label_base else label_base
            records.append(
                SampleRecord(
                    id=new_id("sample"),
                    label=f"{label_base}_{sheet}_{name}",
                    sections={"세특": {subject: content}},
                    source_file=str(path),
                )
            )

    wb.close()
    return records


def parse_xlsx_records(path: Path) -> list[SampleRecord]:
    try:
        records = _parse_xlsx_rows(path)
        if records:
            return records
    except Exception:
        pass

    try:
        text = _xlsx_to_text(path)
    except ImportError as exc:
        raise ValueError("엑셀 처리 패키지(openpyxl)가 설치되지 않았습니다. NAS update 후 재시도하세요.") from exc
    except Exception as exc:
        raise ValueError(f"엑셀 파일을 읽을 수 없습니다: {path.name} ({exc})") from exc

    record = _records_from_document_text(path, text)[0]
    record.label = _label_from_path(path)
    return [record]


def _docx_to_text(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        line = para.text.strip()
        if line:
            parts.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def parse_docx_records(path: Path) -> list[SampleRecord]:
    return _records_from_document_text(path, _docx_to_text(path))
